from flask import Flask, render_template, redirect, url_for, request, flash, Blueprint, jsonify, current_app
from models import db, Usuario, LibroCompra, LibroVenta, Cliente, RetencionVenta, RetencionCompra, IGTFVenta
from datetime import datetime, time, timedelta
from flask_login import login_user, login_required, current_user, logout_user
from functools import wraps
import pytz
import re


main_routes = Blueprint('main', __name__)



# Decorador para verificar si el usuario es admin
def admin_required(f):
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if current_user.rol != 'admin':
            flash('No tienes permiso para acceder a esta página', 'danger')
            return redirect(url_for('main.login'))  # Redirigir al login si no es admin
        return f(*args, **kwargs)
    return decorated_function



@main_routes.route('/', methods=['GET', 'POST'])
def login():
    try:
        if request.method == 'POST':
            nombreuser = request.form.get('nombreuser', '').strip()
            password = request.form.get('password', '').strip()
            
            # Validaciones básicas
            if not nombreuser or not password:
                flash('Todos los campos son obligatorios', 'danger')
                return redirect(url_for('main.login'))
            
            # Buscar al usuario por nombre de usuario
            usuario = db.session.query(Usuario).filter_by(nombreuser=nombreuser).first()
            
            if not usuario:
                flash('Este usuario no es válido', 'danger')  # Mensaje genérico por seguridad
                return redirect(url_for('main.login'))
            
            # Verificar si es admin
            if usuario.rol != 'admin':
                flash('Acceso restringido a solo administradores', 'danger')
                return redirect(url_for('main.login'))
            
            # Verificar contraseña
            if usuario.check_password(password):
                login_user(usuario)
                flash('Inicio de sesión exitoso', 'success')
                return redirect(url_for('main.clientes'))
            else:
                flash('Contraseña inválida', 'danger')
                
    except Exception as e:
        current_app.logger.error(f'Error en login: {str(e)}')
        flash('Ocurrió un error inesperado. Por favor, inténtalo de nuevo más tarde.', 'danger')
    
    return render_template('login.html')



@main_routes.route('/clientes', methods=['GET', 'POST'])
@admin_required
def clientes():
    try:
        clientes = Cliente.query.all()
        return render_template('clientes.html', clientes=clientes)
    except Exception as e:
        current_app.logger.error(f'Error al cargar clientes: {str(e)}')
        flash('Error al cargar la lista de clientes', 'danger')
        return redirect(url_for('main.login'))

@main_routes.route('/logout')
@admin_required
def logout():
    logout_user()  # Cierra la sesión
    flash('Has cerrado sesión exitosamente', 'success')
    return redirect(url_for('main.login'))

@main_routes.route('/usuarios')
@admin_required
def usuarios():
    usuarios = Usuario.query.all()
    return render_template('usuarios.html', usuarios=usuarios)

@main_routes.route('/toggle_block/<int:user_id>', methods=['POST'])
@admin_required
def toggle_block(user_id):
    try:
        usuario = Usuario.query.filter_by(usuario_id=user_id).first_or_404()
        # No permitir que un admin se desactive a sí mismo
        if usuario.usuario_id == current_user.usuario_id:
            return jsonify({'success': False, 'message': 'No puedes desactivarte a ti mismo'}), 400
        
        usuario.status = not usuario.status
        db.session.commit()
        
        return jsonify({'success': True}), 200
    except Exception as e:
        current_app.logger.error(f'Error al cambiar estado del usuario: {str(e)}')
        return jsonify({'success': False, 'message': 'Error al procesar la solicitud'}), 500

@main_routes.route('/librocompra')
@admin_required
def librocompra():
    try:
        compras = LibroCompra.query.all()
        return render_template('librocompra.html', compras=compras)
    except Exception as e:
        current_app.logger.error(f'Error al cargar libro de compras: {str(e)}')
        flash('Error al cargar el libro de compras', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/libroventa')
@admin_required
def libroventa():
    try:
        ventas = LibroVenta.query.all()
        try:
            igtf_ventas = {i.idfacturaventa: i for i in IGTFVenta.query.all()}
        except Exception:
            igtf_ventas = {}
        return render_template('libroventa.html', ventas=ventas, igtf_ventas=igtf_ventas)
    except Exception as e:
        current_app.logger.error(f'Error al cargar libro de ventas: {str(e)}')
        flash('Error al cargar el libro de ventas', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/editar_venta/<int:venta_id>', methods=['GET', 'POST'])
@admin_required
def editar_venta(venta_id):
    venta = LibroVenta.query.get_or_404(venta_id)
    if request.method == 'POST':
        try:
            venta.numerofactura = request.form['numerofactura']
            venta.numerocontrol = request.form.get('numerocontrol', '')
            venta.fechafactura = datetime.strptime(request.form['fechafactura'], '%Y-%m-%d')
            venta.rif = request.form['rif']
            venta.cliente = request.form['cliente']
            venta.montoTotal = float(request.form['montoTotal'])
            venta.base = float(request.form['base'])
            venta.iva = float(request.form['iva'])
            venta.exentas = float(request.form['exentas']) if request.form.get('exentas') else 0.00
            venta.porcentaje_iva = float(request.form['porcentaje_iva'])
            db.session.commit()

            # Lógica IGTF
            monto_igtf = request.form.get('monto_igtf')
            porcentaje_igtf = request.form.get('porcentaje_igtf')
            tasa = request.form.get('tasa')
            cantidad_dolares = request.form.get('cantidad_dolares')
            igtf = IGTFVenta.query.filter_by(idfacturaventa=venta.idfacturaventa).first()
            if any([monto_igtf, porcentaje_igtf, tasa, cantidad_dolares]):
                if igtf:
                    igtf.monto_igtf = monto_igtf or 0
                    igtf.porcentaje_igtf = porcentaje_igtf or 0
                    igtf.tasa = tasa or 0
                    igtf.cantidad_dolares = cantidad_dolares or 0
                else:
                    igtf = IGTFVenta(
                        idfacturaventa=venta.idfacturaventa,
                        monto_igtf=monto_igtf or 0,
                        porcentaje_igtf=porcentaje_igtf or 0,
                        tasa=tasa or 0,
                        cantidad_dolares=cantidad_dolares or 0
                    )
                    db.session.add(igtf)
            elif igtf:
                db.session.delete(igtf)
            db.session.commit()

            flash('Venta actualizada con éxito', 'success')
            return redirect(url_for('main.cliente_libro_ventas', cliente_id=venta.id_cliente))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error al editar venta: {str(e)}')
            flash('Error al editar la venta', 'danger')
    return render_template('editar_venta.html', venta=venta)

@main_routes.route('/editar_compra/<int:compra_id>', methods=['GET', 'POST'])
@admin_required
def editar_compra(compra_id):
    compra = LibroCompra.query.get_or_404(compra_id)
    if request.method == 'POST':
        try:
            compra.numerofactura = request.form['numerofactura']
            compra.control = request.form.get('control', '')
            compra.fechafactura = datetime.strptime(request.form['fechafactura'], '%Y-%m-%d')
            compra.rif = request.form['rif_proveedor']
            compra.provedor = request.form['provedor']
            compra.montoTotal = float(request.form['montoTotal'])
            compra.base = float(request.form['base'])
            compra.iva = float(request.form['iva'])
            compra.exentas = float(request.form['exentas']) if request.form.get('exentas') else 0.00
            compra.facturapolar = True if request.form.get('facturapolar') else False
            compra.documento = request.form.get('documento', 'Factura')
            compra.porcentaje_iva = float(request.form['porcentaje_iva'])
            
            db.session.commit()
            flash('Compra actualizada con éxito', 'success')
            return redirect(url_for('main.cliente_libro_compras', cliente_id=compra.id_cliente))
        except Exception as e:
            db.session.rollback()
            current_app.logger.error(f'Error al editar compra: {str(e)}')
            flash('Error al editar la compra', 'danger')
    return render_template('editar_compra.html', compra=compra)

@main_routes.route('/eliminar_compra/<int:compra_id>', methods=['POST'])
@admin_required
def eliminar_compra(compra_id):
    try:
        compra = LibroCompra.query.get_or_404(compra_id)
        db.session.delete(compra)
        db.session.commit()
        flash('Compra eliminada con éxito', 'success')
    except Exception as e:
        current_app.logger.error(f'Error al eliminar compra: {str(e)}')
        flash('Error al eliminar la compra', 'danger')
    return redirect(url_for('main.librocompra'))

@main_routes.route('/eliminar_venta/<int:venta_id>', methods=['POST'])
@admin_required
def eliminar_venta(venta_id):
    try:
        venta = LibroVenta.query.get_or_404(venta_id)
        db.session.delete(venta)
        db.session.commit()
        flash('Venta eliminada con éxito', 'success')
    except Exception as e:
        current_app.logger.error(f'Error al eliminar venta: {str(e)}')
        flash('Error al eliminar la venta', 'danger')
    return redirect(url_for('main.libroventa'))

@main_routes.route('/toggle_status/<int:cliente_id>', methods=['POST'])
@admin_required
def toggle_status(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        # Cambiar el estado entre 'activo' e 'inactivo'
        cliente.status = 'inactivo' if cliente.status == 'activo' else 'activo'
        db.session.commit()
        return jsonify({'success': True}), 200
    except Exception as e:
        current_app.logger.error(f'Error al cambiar estado del cliente: {str(e)}')
        return jsonify({'success': False, 'message': 'Error al procesar la solicitud'}), 500

@main_routes.route('/cliente/<int:cliente_id>/libro-ventas')
@admin_required
def cliente_libro_ventas(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        ventas = LibroVenta.query.filter_by(id_cliente=cliente_id).all()
        retenciones = {r.idfacturaventa: r for r in RetencionVenta.query.filter(
            RetencionVenta.idfacturaventa.in_([v.idfacturaventa for v in ventas])
        ).all()}
        try:
            igtf_ventas = {i.idfacturaventa: i for i in IGTFVenta.query.filter(
                IGTFVenta.idfacturaventa.in_([v.idfacturaventa for v in ventas])
            ).all()}
        except Exception:
            igtf_ventas = {}
        return render_template('libroventa.html', 
                             ventas=ventas, 
                             cliente=cliente, 
                             filtro_cliente=True,
                             retenciones=retenciones,
                             igtf_ventas=igtf_ventas)
    except Exception as e:
        current_app.logger.error(f'Error al cargar libro de ventas del cliente: {str(e)}')
        flash('Error al cargar el libro de ventas del cliente', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/cliente/<int:cliente_id>/libro-compras')
@admin_required
def cliente_libro_compras(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        compras = LibroCompra.query.filter_by(id_cliente=cliente_id).all()
        
        # Obtener las retenciones existentes para cada factura
        retenciones = {r.idfacturacompra: r for r in RetencionCompra.query.filter(
            RetencionCompra.idfacturacompra.in_([c.idfacturacompra for c in compras])
        ).all()}
        
        return render_template('librocompra.html', 
                             compras=compras, 
                             cliente=cliente, 
                             filtro_cliente=True,
                             retenciones=retenciones)
    except Exception as e:
        current_app.logger.error(f'Error al cargar libro de compras del cliente: {str(e)}')
        flash('Error al cargar el libro de compras del cliente', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/cliente/<int:cliente_id>')
@admin_required
def ver_cliente(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        return render_template('ver_cliente.html', cliente=cliente)
    except Exception as e:
        current_app.logger.error(f'Error al cargar detalles del cliente: {str(e)}')
        flash('Error al cargar los detalles del cliente', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/crear_cliente', methods=['GET'])
@admin_required
def mostrar_crear_cliente():
    try:
        return render_template('crear_cliente.html')
    except Exception as e:
        current_app.logger.error(f'Error al mostrar formulario de cliente: {str(e)}')
        flash('Error al cargar el formulario', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/crear_cliente', methods=['POST'])
@admin_required
def crear_cliente():
    try:
        # Obtener datos del formulario
        nombre = request.form.get('nombre')
        representante = request.form.get('representante')
        rif = request.form.get('rif')
        nit = request.form.get('nit')
        direccion = request.form.get('direccion')
        telefono = request.form.get('telefono')
        numero = request.form.get('numero')
        contribuyente = request.form.get('contribuyente')  # Ahora será 'ordinario' o 'especial'
        registradora = request.form.get('registradora')
        clasificacion = request.form.get('clasificacion')
        usuarioseniat = request.form.get('usuarioseniat')
        claveseniat = request.form.get('claveseniat')
        clavepatente = request.form.get('clavepatente')
        directoriocomo = request.form.get('directoriocomo')
        
        # Validar campos requeridos
        if not nombre:
            flash('El nombre del cliente es requerido', 'danger')
            return redirect(url_for('main.mostrar_crear_cliente'))
            
        if contribuyente not in ['ordinario', 'especial']:
            flash('El tipo de contribuyente debe ser ordinario o especial', 'danger')
            return redirect(url_for('main.mostrar_crear_cliente'))
            
        # Crear nuevo cliente
        nuevo_cliente = Cliente(
            nombre=nombre,
            representante=representante,
            rif=rif,
            nit=nit,
            direccion=direccion,
            telefono=telefono,
            numero=numero,
            contribuyente=contribuyente,
            registradora=registradora,
            clasificacion=clasificacion,
            usuarioseniat=usuarioseniat,
            claveseniat=claveseniat,
            clavepatente=clavepatente,
            directoriocomo=directoriocomo,
            status='activo'
        )
        
        db.session.add(nuevo_cliente)
        db.session.commit()
        
        flash('Cliente creado exitosamente', 'success')
        return redirect(url_for('main.clientes'))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error al crear cliente: {str(e)}')
        flash('Error al crear el cliente', 'danger')
        return redirect(url_for('main.mostrar_crear_cliente'))

@main_routes.route('/crear-venta', methods=['GET', 'POST'])
@admin_required
def crear_venta():
    cliente_id = request.args.get('cliente_id')
    if not cliente_id:
        flash('Se requiere especificar un cliente', 'danger')
        return redirect(url_for('main.clientes'))
        
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        
        if request.method == 'POST':
            try:
                nueva_venta = LibroVenta(
                    id_cliente=int(cliente_id),
                    numerofactura=request.form['numerofactura'],
                    fechafactura=datetime.strptime(request.form['fechafactura'], '%Y-%m-%d'),
                    rif=request.form['rif'],
                    cliente=request.form['cliente'],
                    montoTotal=float(request.form['montoTotal']),
                    base=float(request.form['base']),
                    iva=float(request.form['iva']),
                    exentas=float(request.form['exentas']) if request.form.get('exentas') else 0.00,
                    documento='Factura',
                    numerocontrol=request.form.get('numerocontrol', '')
                )
                
                db.session.add(nueva_venta)
                db.session.commit()
                
                # Después de guardar la venta, guardar IGTF
                monto_igtf = request.form.get('monto_igtf')
                porcentaje_igtf = request.form.get('porcentaje_igtf')
                tasa = request.form.get('tasa')
                cantidad_dolares = request.form.get('cantidad_dolares')

                if any([monto_igtf, porcentaje_igtf, tasa, cantidad_dolares]):
                    igtf = IGTFVenta(
                        idfacturaventa=nueva_venta.idfacturaventa,
                        monto_igtf=monto_igtf or 0,
                        porcentaje_igtf=porcentaje_igtf or 0,
                        tasa=tasa or 0,
                        cantidad_dolares=cantidad_dolares or 0
                    )
                    db.session.add(igtf)
                    db.session.commit()
                
                flash('Venta creada exitosamente', 'success')
                return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
                
            except Exception as e:
                db.session.rollback()
                flash(f'Error al crear la venta: {str(e)}', 'danger')
                return render_template('crear_venta.html', cliente=cliente)
        
        # GET request
        return render_template('crear_venta.html', cliente=cliente)
        
    except Exception as e:
        flash('Error al acceder al cliente', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/crear-compra', methods=['GET', 'POST'])
@admin_required
def crear_compra():
    cliente_id = request.args.get('cliente_id')
    if not cliente_id:
        flash('Se requiere especificar un cliente', 'danger')
        return redirect(url_for('main.clientes'))
        
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        
        if request.method == 'POST':
            try:
                nueva_compra = LibroCompra(
                    id_cliente=int(cliente_id),
                    numerofactura=request.form['numerofactura'],
                    control=request.form['control'],
                    fechafactura=datetime.strptime(request.form['fechafactura'], '%Y-%m-%d'),
                    rif=request.form['rif_proveedor'],
                    provedor=request.form['provedor'],
                    montoTotal=float(request.form['montoTotal']),
                    base=float(request.form['base']),
                    iva=float(request.form['iva']),
                    exentas=float(request.form['exentas']) if request.form.get('exentas') else 0.00,
                    facturapolar=True if request.form.get('facturapolar') else False,
                    documento=request.form.get('documento', 'Factura')
                )
                
                db.session.add(nueva_compra)
                db.session.commit()
                
                flash('Compra creada exitosamente', 'success')
                return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))
                
            except Exception as e:
                db.session.rollback()
                flash(f'Error al crear la compra: {str(e)}', 'danger')
                return render_template('crear_compra.html', cliente=cliente)
        
        # GET request
        return render_template('crear_compra.html', cliente=cliente)
        
    except Exception as e:
        flash('Error al acceder al cliente', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/editar_cliente/<int:cliente_id>', methods=['GET'])
@admin_required
def mostrar_editar_cliente(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        return render_template('editar_cliente.html', cliente=cliente)
    except Exception as e:
        current_app.logger.error(f'Error al cargar formulario de edición: {str(e)}')
        flash('Error al cargar el formulario de edición', 'danger')
        return redirect(url_for('main.clientes'))

@main_routes.route('/editar_cliente/<int:cliente_id>', methods=['POST'])
@admin_required
def editar_cliente(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        
        # Actualizar datos del cliente
        cliente.nombre = request.form.get('nombre')
        cliente.representante = request.form.get('representante')
        cliente.rif = request.form.get('rif')
        cliente.nit = request.form.get('nit')
        cliente.direccion = request.form.get('direccion')
        cliente.telefono = request.form.get('telefono')
        cliente.numero = request.form.get('numero')
        cliente.contribuyente = request.form.get('contribuyente')
        cliente.registradora = request.form.get('registradora')
        cliente.clasificacion = request.form.get('clasificacion')
        cliente.usuarioseniat = request.form.get('usuarioseniat')
        
        # Solo actualizar la clave SENIAT si se proporciona una nueva
        nueva_clave_seniat = request.form.get('claveseniat')
        if nueva_clave_seniat:
            cliente.claveseniat = nueva_clave_seniat
            
        cliente.clavepatente = request.form.get('clavepatente')
        cliente.directoriocomo = request.form.get('directoriocomo')
        
        db.session.commit()
        flash('Cliente actualizado exitosamente', 'success')
        return redirect(url_for('main.ver_cliente', cliente_id=cliente_id))
        
    except Exception as e:
        db.session.rollback()
        current_app.logger.error(f'Error al actualizar cliente: {str(e)}')
        flash('Error al actualizar el cliente', 'danger')
        return redirect(url_for('main.mostrar_editar_cliente', cliente_id=cliente_id))

from flask import request, jsonify, send_from_directory
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup
import os
from datetime import datetime
import re

def set_cell_border(cell, **kwargs):
    """Establecer bordes de celda"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    sides = ['top', 'left', 'bottom', 'right']
    for side in sides:
        if side in kwargs:
            tag = f'w:{side}'
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            
            element.set(qn('w:val'), kwargs[side].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[side].get('sz', 4)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), kwargs[side].get('color', '000000'))

def format_number_spanish(num):
    """Formatear número con separadores españoles"""
    try:
        if num == 0:
            return "0.00"
        return f"{float(num):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except:
        return "0.00"

def parse_number_spanish(num_str):
    if not num_str:
        return 0.0
    return float(num_str.replace('.', '').replace(',', '.'))

@main_routes.route('/exportar_libro_compras_word', methods=['POST'])
def exportar_libro_compras_word():
    try:
        data = request.get_json()
        html_content = data.get('html_content')
        cliente = data.get('cliente', {})
        
        # Asegurar que el directorio temp existe
        os.makedirs('temp', exist_ok=True)
        
        # Parsear HTML para extraer datos
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Crear documento Word
        doc = Document()
        
        # Configurar márgenes ajustados
        for section in doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(0.8)
            section.right_margin = Cm(0.8)
            # Configurar orientación horizontal si es necesario
            section.page_width = Cm(29.7)  # A4 horizontal
            section.page_height = Cm(21.0)
        
        # Configurar estilo base - Courier New para alineación perfecta
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(8)
        
        # === ENCABEZADO SUPERIOR ===
        # Crear tabla para el encabezado (2 columnas)
        header_table = doc.add_table(rows=1, cols=2)
        header_table.columns[0].width = Cm(12)
        header_table.columns[1].width = Cm(15)

        # Título
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(f"{cliente.get('nombre', 'N/A')}\n")
        title_run.bold = True
        title_run.font.size = Pt(9)
        
        left_para = header_table.rows[0].cells[0].paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_para.add_run(f"R.I.F.  : {cliente.get('rif', 'N/A')}\n")
        left_para.add_run(f"Direccion : {cliente.get('direccion', 'N/A')}\n")

        
        # Extraer período
        period_text = "ABRIL - 2025  periodo desde 01/04/2025 hasta 30/04/2025"
        company_info = soup.find('div', class_='company-info')
        if company_info:
            for text in company_info.stripped_strings:
                if "Mes :" in text:
                    period_text = text.replace("Mes : ", "")
                    break
        
        left_para.add_run(f"Mes : {period_text}\n")
        left_para.add_run("Pag:1")
        

        
        # Quitar bordes de la tabla de encabezado
        for row in header_table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={'val': 'nil'}, bottom={'val': 'nil'}, 
                              left={'val': 'nil'}, right={'val': 'nil'})
        
        # Espacio
        doc.add_paragraph()
        
        # === TABLA PRINCIPAL DE DATOS ===
        # Crear tabla con 15 columnas (como en la imagen)
        main_table = doc.add_table(rows=4, cols=15)
        
        # Configurar anchos de columnas para coincidir exactamente con la imagen
        column_widths = [
            Cm(1.2),  # FECHA
            Cm(1.2),  # FACTURA FECHA  
            Cm(1.2),  # FACTURA NUMERO
            Cm(0.8),  # DOC
            Cm(1.5),  # CONTROL
            Cm(4.0),  # NOMBRE DEL PROVEEDOR
            Cm(2.2),  # R.I.F.
            Cm(1.8),  # TOTAL COMPRAS INCLUYENDO I.V.A.
            Cm(1.5),  # SIN DERECHO A CREDITO
            Cm(1.2),  # BASE (IMPORTACION)
            Cm(1.2),  # I.V.A. (IMPORTACION)
            Cm(1.2),  # BASE (NACIONAL)
            Cm(1.2),  # I.V.A. (NACIONAL)
            Cm(1.5),  # Comprobante
            Cm(1.2)   # RETENCION
        ]
        
        for i, width in enumerate(column_widths):
            for row in main_table.rows:
                row.cells[i].width = width
        
        # === PRIMERA FILA - ENCABEZADOS PRINCIPALES ===
        row1 = main_table.rows[0]
        
        # Combinar celdas para "DATOS DE LAS FACTURA" (columnas 0-4)
        merge_cell1 = row1.cells[0]
        for i in range(1, 5):
            merge_cell1.merge(row1.cells[i])
        merge_cell1.text = "DATOS DE LAS FACTURA"
        para1 = merge_cell1.paragraphs[0]
        para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = para1.runs[0]
        run1.bold = True
        run1.font.size = Pt(8)
        
        # NOMBRE DEL PROVEEDOR
        row1.cells[5].text = ""
        
        # R.I.F.
        row1.cells[6].text = ""
        
        # TOTAL COMPRAS INCLUYENDO I.V.A.
        row1.cells[7].text = "TOTAL COMPRAS\nINCLUYENDO\nI.V.A."
        para7 = row1.cells[7].paragraphs[0]
        para7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para7.runs[0].bold = True
        para7.runs[0].font.size = Pt(7)
        
        # SIN DERECHO A CREDITO
        row1.cells[8].text = "SIN\nDERECHO\nA\nCREDITO"
        para8 = row1.cells[8].paragraphs[0]
        para8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para8.runs[0].bold = True
        para8.runs[0].font.size = Pt(7)
        
        # Combinar "CON DERECHO A CREDITO" (columnas 9-12)
        merge_cell2 = row1.cells[9]
        for i in range(10, 13):
            merge_cell2.merge(row1.cells[i])
        merge_cell2.text = "CON DERECHO A CREDITO"
        para9 = merge_cell2.paragraphs[0]
        para9.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para9.runs[0].bold = True
        para9.runs[0].font.size = Pt(8)
        
        # Celdas finales vacías
        row1.cells[13].text = ""
        row1.cells[14].text = ""
        
        # === SEGUNDA FILA - SUBENCABEZADOS ===
        row2 = main_table.rows[1]
        headers2 = [
            "", "", "", "", "",
            "", "", "",
            "", "", "IMPORTACION", "", "NACIONAL", "", ""
        ]
        
        for i, header in enumerate(headers2):
            if header:
                row2.cells[i].text = header
                para = row2.cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(7)
        # Quitar el borde inferior de la segunda fila de encabezados
        for cell in row2.cells:
            set_cell_border(cell, bottom={'val': 'nil'})
        
        # === TERCERA FILA - ENCABEZADOS DETALLADOS ===
        row3 = main_table.rows[2]
        headers3 = [
            "FECHA", "FECHA FACTURA", "FACTURA NUMERO", "DOC", "CONTROL",
            "NOMBRE DEL PROVEEDOR", "R.I.F.", "", "", "BASE", "I.V.A.", "BASE", "I.V.A.", "Comprobante", "RETENCION"
        ]
        
        for i, header in enumerate(headers3):
            if header:
                row3.cells[i].text = header
                para = row3.cells[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].bold = True
                para.runs[0].font.size = Pt(7)
        # Quitar el borde inferior de la tercera fila de encabezados
        for cell in row3.cells:
            set_cell_border(cell, bottom={'val': 'nil'})
        
        # === EXTRAER Y AGREGAR DATOS ===
        html_table = soup.find('table')
        totals = {
            'totalCompras': 0, 'totalExentos': 0, 'totalBaseImportacion': 0,
            'totalIvaImportacion': 0, 'totalBaseNacional': 0, 'totalIvaNacional': 0,
            'totalRetencion': 0
        }
        
        if html_table:
            for row in html_table.find_all('tr'):
                cells = row.find_all('td')
                if len(cells) >= 14 and not any(cls in row.get('class', []) for cls in ['subtotal', 'total']):
                    # Agregar fila de datos
                    data_row = main_table.add_row()
                    
                    # Mapear datos según la estructura esperada
                    cell_data = []
                    for i, cell in enumerate(cells[:15]):
                        text = cell.get_text().strip()
                        cell_data.append(text)
                        
                        # Acumular totales usando parse_number_spanish
                        if i >= 7 and text:
                            num_value = parse_number_spanish(text)
                            if i == 7: totals['totalCompras'] += num_value
                            elif i == 8: totals['totalExentos'] += num_value
                            elif i == 9: totals['totalBaseImportacion'] += num_value
                            elif i == 10: totals['totalIvaImportacion'] += num_value
                            elif i == 11: totals['totalBaseNacional'] += num_value
                            elif i == 12: totals['totalIvaNacional'] += num_value
                            elif i == 14: totals['totalRetencion'] += num_value
                    
                    # Llenar celdas de la fila
                    for i, data in enumerate(cell_data):
                        if i < len(data_row.cells):
                            cell = data_row.cells[i]
                            cell.text = data
                            para = cell.paragraphs[0]
                            run = para.runs[0] if para.runs else para.add_run(data)
                            run.font.size = Pt(7)
                            
                            # Alineación
                            if i >= 7:  # Columnas numéricas
                                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # === AGREGAR FILAS DE TOTALES ===
        # Subtotales
        subtotal_row = main_table.add_row()
        subtotal_data = [
            "", "", "", "", "", "", "SUB-TOTALES.",
            format_number_spanish(totals['totalCompras']),
            format_number_spanish(totals['totalExentos']),
            format_number_spanish(totals['totalBaseImportacion']),
            format_number_spanish(totals['totalIvaImportacion']),
            format_number_spanish(totals['totalBaseNacional']),
            format_number_spanish(totals['totalIvaNacional']),
            "",
            format_number_spanish(totals['totalRetencion'])
        ]
        
        for i, data in enumerate(subtotal_data):
            if i < len(subtotal_row.cells):
                cell = subtotal_row.cells[i]
                cell.text = data
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(data)
                run.font.size = Pt(8)
                run.bold = True
                
                if i >= 6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Totales finales
        total_row = main_table.add_row()
        total_data = [
            "", "", "", "", "", "", "TOTALES.",
            format_number_spanish(totals['totalCompras']),
            format_number_spanish(totals['totalExentos']),
            format_number_spanish(totals['totalBaseImportacion']),
            format_number_spanish(totals['totalIvaImportacion']),
            format_number_spanish(totals['totalBaseNacional']),
            format_number_spanish(totals['totalIvaNacional']),
            "",
            format_number_spanish(totals['totalRetencion'])
        ]
        
        for i, data in enumerate(total_data):
            if i < len(total_row.cells):
                cell = total_row.cells[i]
                cell.text = data
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(data)
                run.font.size = Pt(8)
                run.bold = True
                
                if i >= 6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # === APLICAR BORDES ===
        border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
        
        # Solo bordes horizontales en la parte superior e inferior
        for i, row in enumerate(main_table.rows):
            for j, cell in enumerate(row.cells):
                if i <= 3:  # Filas de encabezado
                    set_cell_border(cell, top=border_style, bottom=border_style)
                else:  # Filas de datos
                    if i == len(main_table.rows) - 2:  # Fila de subtotales
                        set_cell_border(cell, top=border_style)
                    elif i == len(main_table.rows) - 1:  # Fila de totales
                        set_cell_border(cell, bottom=border_style)
        
        # Nombre de archivo
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        cliente_name = re.sub(r'[^a-zA-Z0-9_-]', '_', cliente.get('nombre', 'cliente'))
        filename = f"libro_compras_{cliente_name}_{timestamp}.docx"
        filepath = os.path.join('temp', filename)
        
        # Guardar documento
        doc.save(filepath)
        
        return jsonify({'filename': filename})
    
    except Exception as e:
        current_app.logger.error(f'Error al generar Word: {str(e)}')
        return jsonify({'error': f'Error interno: {str(e)}'}, 500)

@main_routes.route('/descargar_word/<filename>')
def descargar_word(filename):
    try:
        return send_from_directory('temp', filename, as_attachment=True)
    except Exception as e:
        return jsonify({'error': f'Archivo no encontrado: {str(e)}'}, 404)

import glob
import time

def limpiar_temporales():
    """Eliminar archivos temporales con más de 1 hora de antigüedad"""
    try:
        for filepath in glob.glob('temp/*.docx'):
            if os.path.getmtime(filepath) < time.time() - 3600:
                os.remove(filepath)
    except Exception as e:
        print(f"Error limpiando temporales: {e}")

# Ejecutar al inicio de la app
limpiar_temporales()

@main_routes.route('/cliente/<int:cliente_id>/retenciones-ventas')
@admin_required
def retenciones_ventas(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        # Permitir retenciones de ventas para todos los clientes
        # if cliente.contribuyente != 'especial':
        #     flash('Este cliente no puede tener retenciones de ventas', 'warning')
        #     return redirect(url_for('main.ver_cliente', cliente_id=cliente_id))
        
        # Obtener todas las retenciones de venta del cliente
        retenciones = db.session.query(RetencionVenta, LibroVenta)\
            .join(LibroVenta, RetencionVenta.idfacturaventa == LibroVenta.idfacturaventa)\
            .filter(LibroVenta.id_cliente == cliente_id)\
            .all()
            
        return render_template('retenciones/lista_retenciones_ventas.html', 
                             cliente=cliente,
                             retenciones=[r[0] for r in retenciones],
                             facturas={r[0].idfacturaventa: r[1] for r in retenciones})
    except Exception as e:
        current_app.logger.error(f'Error al cargar retenciones de ventas: {str(e)}')
        flash('Error al cargar las retenciones de ventas', 'danger')
        return redirect(url_for('main.ver_cliente', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/retenciones-compras')
@admin_required
def retenciones_compras(cliente_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        
        # Bloquear retenciones de compras para clientes ordinarios
        if cliente.contribuyente != 'especial':
            flash('Solo los clientes especiales pueden tener retenciones de compras', 'warning')
            return redirect(url_for('main.ver_cliente', cliente_id=cliente_id))
            
        # Obtener todas las retenciones de compra del cliente
        retenciones = db.session.query(RetencionCompra, LibroCompra)\
            .join(LibroCompra, RetencionCompra.idfacturacompra == LibroCompra.idfacturacompra)\
            .filter(LibroCompra.id_cliente == cliente_id)\
            .all()
            
        return render_template('retenciones/lista_retenciones_compras.html', 
                             cliente=cliente,
                             retenciones=[r[0] for r in retenciones],
                             facturas={r[0].idfacturacompra: r[1] for r in retenciones})
    except Exception as e:
        current_app.logger.error(f'Error al cargar retenciones de compras: {str(e)}')
        flash('Error al cargar las retenciones de compras', 'danger')
        return redirect(url_for('main.ver_cliente', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/ver-retenciones-ventas/<int:factura_id>')
@admin_required
def ver_retenciones_ventas(cliente_id, factura_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        # Permitir retenciones de ventas para todos los clientes
        # if cliente.contribuyente != 'especial':
        #     flash('Este cliente no puede tener retenciones de ventas', 'warning')
        #     return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
            
        factura = LibroVenta.query.get_or_404(factura_id)
        retenciones = RetencionVenta.query.filter_by(idfacturaventa=factura_id).all()
        return render_template('retenciones/lista_retenciones_ventas.html', 
                             cliente=cliente, 
                             factura=factura, 
                             retenciones=retenciones)
    except Exception as e:
        current_app.logger.error(f'Error al cargar retenciones de venta: {str(e)}')
        flash('Error al cargar las retenciones', 'danger')
        return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/ver-retenciones-compras/<int:factura_id>')
@admin_required
def ver_retenciones_compras(cliente_id, factura_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        factura = LibroCompra.query.get_or_404(factura_id)
        retenciones = RetencionCompra.query.filter_by(idfacturacompra=factura_id).all()
        return render_template('retenciones/lista_retenciones_compras.html', 
                             cliente=cliente, 
                             factura=factura, 
                             retenciones=retenciones)
    except Exception as e:
        current_app.logger.error(f'Error al cargar retenciones de compra: {str(e)}')
        flash('Error al cargar las retenciones', 'danger')
        return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/crear-retencion-venta/<int:factura_id>', methods=['GET', 'POST'])
@admin_required
def crear_retencion_venta(cliente_id, factura_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        # Permitir retenciones de ventas para todos los clientes
        # if cliente.contribuyente != 'especial':
        #     flash('Este cliente no puede tener retenciones de ventas', 'warning')
        #     return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
            
        factura = LibroVenta.query.get_or_404(factura_id)
        
        # Verificar si ya existe una retención para esta factura
        retencion_existente = RetencionVenta.query.filter_by(idfacturaventa=factura_id).first()
        if retencion_existente:
            flash('Ya existe una retención para esta factura', 'warning')
            return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
        
        if request.method == 'POST':
            try:
                # Obtener y validar los datos del formulario
                numero_comprobante = request.form.get('numero_comprobante')
                fecha_emision = request.form.get('fecha_emision')
                porcentaje_retencion = request.form.get('porcentaje_retencion')
                valor_retencion = request.form.get('valor_retencion')
                
                # Validar que todos los campos necesarios estén presentes
                if not all([numero_comprobante, fecha_emision, porcentaje_retencion, valor_retencion]):
                    flash('Todos los campos son requeridos', 'danger')
                    return render_template('retenciones/crear_retencion_venta.html', 
                                         cliente=cliente, 
                                         factura=factura)
                
                # Convertir y validar valores numéricos
                try:
                    porcentaje = float(porcentaje_retencion)
                    valor = float(valor_retencion)
                except ValueError:
                    flash('Los valores numéricos no son válidos', 'danger')
                    return render_template('retenciones/crear_retencion_venta.html', 
                                         cliente=cliente, 
                                         factura=factura)
                
                # Crear la retención
                nueva_retencion = RetencionVenta(
                    id_cliente=cliente_id,
                    idfacturaventa=factura_id,
                    numero_comprobante=numero_comprobante,
                    fecha_emision=datetime.strptime(fecha_emision, '%Y-%m-%d'),
                    porcentaje_retencion=porcentaje,
                    valor_retencion=valor
                )
                
                # Imprimir información de depuración
                print(f"Nueva retención a crear: {nueva_retencion}")
                print(f"Datos: cliente_id={cliente_id}, factura_id={factura_id}, numero={numero_comprobante}, fecha={fecha_emision}, porcentaje={porcentaje}, valor={valor}")
                
                db.session.add(nueva_retencion)
                db.session.commit()
                
                # Después de guardar la retención, guardar IGTF
                monto_igtf = request.form.get('monto_igtf')
                porcentaje_igtf = request.form.get('porcentaje_igtf')
                tasa = request.form.get('tasa')
                cantidad_dolares = request.form.get('cantidad_dolares')

                if any([monto_igtf, porcentaje_igtf, tasa, cantidad_dolares]):
                    igtf = IGTFVenta.query.filter_by(idfacturaventa=nueva_retencion.idfacturaventa).first()
                    if igtf:
                        igtf.monto_igtf = monto_igtf or 0
                        igtf.porcentaje_igtf = porcentaje_igtf or 0
                        igtf.tasa = tasa or 0
                        igtf.cantidad_dolares = cantidad_dolares or 0
                    else:
                        igtf = IGTFVenta(
                            idfacturaventa=nueva_retencion.idfacturaventa,
                            monto_igtf=monto_igtf or 0,
                            porcentaje_igtf=porcentaje_igtf or 0,
                            tasa=tasa or 0,
                            cantidad_dolares=cantidad_dolares or 0
                        )
                        db.session.add(igtf)
                db.session.commit()
                
                flash('Retención de venta creada exitosamente', 'success')
                return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
                
            except ValueError as e:
                db.session.rollback()
                print(f"Error de valor: {str(e)}")
                flash(f'Error en los valores ingresados: {str(e)}', 'danger')
                return render_template('retenciones/crear_retencion_venta.html', 
                                     cliente=cliente, 
                                     factura=factura)
            except Exception as e:
                db.session.rollback()
                print(f"Error general: {str(e)}")
                flash(f'Error al crear la retención: {str(e)}', 'danger')
                return render_template('retenciones/crear_retencion_venta.html', 
                                     cliente=cliente, 
                                     factura=factura)
        
        return render_template('retenciones/crear_retencion_venta.html', 
                             cliente=cliente, 
                             factura=factura)
                             
    except Exception as e:
        print(f"Error en la ruta: {str(e)}")
        flash('Error al procesar la solicitud', 'danger')
        return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/crear-retencion-compra/<int:factura_id>', methods=['GET', 'POST'])
@admin_required
def crear_retencion_compra(cliente_id, factura_id):
    try:
        cliente = Cliente.query.get_or_404(cliente_id)
        factura = LibroCompra.query.get_or_404(factura_id)
        
        # Verificar si ya existe una retención para esta factura
        retencion_existente = RetencionCompra.query.filter_by(idfacturacompra=factura_id).first()
        if retencion_existente:
            flash('Ya existe una retención para esta factura', 'warning')
            return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))
        
        if request.method == 'POST':
            try:
                # Obtener y validar los datos del formulario
                numero_comprobante = request.form.get('numero_comprobante')
                fecha_emision = request.form.get('fecha_emision')
                porcentaje_retencion = request.form.get('porcentaje_retencion')
                valor_retencion = request.form.get('valor_retencion')
                
                # Validar que todos los campos necesarios estén presentes
                if not all([numero_comprobante, fecha_emision, porcentaje_retencion, valor_retencion]):
                    flash('Todos los campos son requeridos', 'danger')
                    return render_template('retenciones/crear_retencion_compra.html', 
                                         cliente=cliente, 
                                         factura=factura)
                
                # Convertir y validar valores numéricos
                try:
                    porcentaje = float(porcentaje_retencion)
                    valor = float(valor_retencion)
                except ValueError:
                    flash('Los valores numéricos no son válidos', 'danger')
                    return render_template('retenciones/crear_retencion_compra.html', 
                                         cliente=cliente, 
                                         factura=factura)
                
                # Crear la retención
                nueva_retencion = RetencionCompra(
                    id_cliente=cliente_id,
                    idfacturacompra=factura_id,
                    numero_comprobante=numero_comprobante,
                    fecha_emision=datetime.strptime(fecha_emision, '%Y-%m-%d'),
                    porcentaje_retencion=porcentaje,
                    valor_retencion=valor
                )
                
                # Imprimir información de depuración
                print(f"Nueva retención a crear: {nueva_retencion}")
                print(f"Datos: cliente_id={cliente_id}, factura_id={factura_id}, numero={numero_comprobante}, fecha={fecha_emision}, porcentaje={porcentaje}, valor={valor}")
                
                db.session.add(nueva_retencion)
                db.session.commit()
                
                flash('Retención de compra creada exitosamente', 'success')
                return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))
                
            except ValueError as e:
                db.session.rollback()
                print(f"Error de valor: {str(e)}")
                flash(f'Error en los valores ingresados: {str(e)}', 'danger')
                return render_template('retenciones/crear_retencion_compra.html', 
                                     cliente=cliente, 
                                     factura=factura)
            except Exception as e:
                db.session.rollback()
                print(f"Error general: {str(e)}")
                flash(f'Error al crear la retención: {str(e)}', 'danger')
                return render_template('retenciones/crear_retencion_compra.html', 
                                     cliente=cliente, 
                                     factura=factura)
        
        return render_template('retenciones/crear_retencion_compra.html', 
                             cliente=cliente, 
                             factura=factura)
                             
    except Exception as e:
        print(f"Error en la ruta: {str(e)}")
        flash('Error al procesar la solicitud', 'danger')
        return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/editar-retencion-venta/<int:retencion_id>', methods=['GET', 'POST'])
@admin_required
def editar_retencion_venta(cliente_id, retencion_id):
    try:
        retencion = RetencionVenta.query.get_or_404(retencion_id)
        cliente = Cliente.query.get_or_404(cliente_id)
        factura = LibroVenta.query.get_or_404(retencion.idfacturaventa)
        
        if request.method == 'POST':
            retencion.numero_comprobante = request.form['numero_comprobante']
            retencion.fecha_emision = datetime.strptime(request.form['fecha_emision'], '%Y-%m-%d')
            retencion.porcentaje_retencion = request.form['porcentaje_retencion']
            retencion.valor_retencion = float(request.form['valor_retencion'])
            
            db.session.commit()
            flash('Retención de venta actualizada exitosamente', 'success')
            return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))
        
        return render_template('retenciones/editar_retencion_venta.html',
                             cliente=cliente,
                             factura=factura,
                             retencion=retencion)
    except Exception as e:
        current_app.logger.error(f'Error al editar retención de venta: {str(e)}')
        flash('Error al editar la retención', 'danger')
        return redirect(url_for('main.cliente_libro_ventas', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/editar-retencion-compra/<int:retencion_id>', methods=['GET', 'POST'])
@admin_required
def editar_retencion_compra(cliente_id, retencion_id):
    try:
        retencion = RetencionCompra.query.get_or_404(retencion_id)
        cliente = Cliente.query.get_or_404(cliente_id)
        factura = LibroCompra.query.get_or_404(retencion.idfacturacompra)
        
        if request.method == 'POST':
            retencion.numero_comprobante = request.form['numero_comprobante']
            retencion.fecha_emision = datetime.strptime(request.form['fecha_emision'], '%Y-%m-%d')
            retencion.porcentaje_retencion = request.form['porcentaje_retencion']
            retencion.valor_retencion = float(request.form['valor_retencion'])
            
            db.session.commit()
            flash('Retención de compra actualizada exitosamente', 'success')
            return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))
        
        return render_template('retenciones/editar_retencion_compra.html',
                             cliente=cliente,
                             factura=factura,
                             retencion=retencion)
    except Exception as e:
        current_app.logger.error(f'Error al editar retención de compra: {str(e)}')
        flash('Error al editar la retención', 'danger')
        return redirect(url_for('main.cliente_libro_compras', cliente_id=cliente_id))

@main_routes.route('/cliente/<int:cliente_id>/eliminar-retencion-venta/<int:retencion_id>', methods=['POST'])
@admin_required
def eliminar_retencion_venta(cliente_id, retencion_id):
    try:
        retencion = RetencionVenta.query.get_or_404(retencion_id)
        factura_id = retencion.idfacturaventa
        
        db.session.delete(retencion)
        db.session.commit()
        
        flash('Retención de venta eliminada exitosamente', 'success')
        return redirect(url_for('main.ver_retenciones_ventas', 
                              cliente_id=cliente_id, 
                              factura_id=factura_id))
    except Exception as e:
        current_app.logger.error(f'Error al eliminar retención de venta: {str(e)}')
        flash('Error al eliminar la retención', 'danger')
        return redirect(url_for('main.ver_retenciones_ventas', 
                              cliente_id=cliente_id, 
                              factura_id=factura_id))

@main_routes.route('/cliente/<int:cliente_id>/eliminar-retencion-compra/<int:retencion_id>', methods=['POST'])
@admin_required
def eliminar_retencion_compra(cliente_id, retencion_id):
    try:
        retencion = RetencionCompra.query.get_or_404(retencion_id)
        factura_id = retencion.idfacturacompra
        
        db.session.delete(retencion)
        db.session.commit()
        
        flash('Retención de compra eliminada exitosamente', 'success')
        return redirect(url_for('main.ver_retenciones_compras', 
                              cliente_id=cliente_id, 
                              factura_id=factura_id))
    except Exception as e:
        current_app.logger.error(f'Error al eliminar retención de compra: {str(e)}')
        flash('Error al eliminar la retención', 'danger')
        return redirect(url_for('main.ver_retenciones_compras', 
                              cliente_id=cliente_id, 
                              factura_id=factura_id))

@main_routes.route('/cliente/<int:cliente_id>/ventas_igtf')
@admin_required
def ventas_igtf(cliente_id):
    cliente = Cliente.query.get_or_404(cliente_id)
    facturas_igtf = db.session.query(LibroVenta, IGTFVenta)\
        .join(IGTFVenta, LibroVenta.idfacturaventa == IGTFVenta.idfacturaventa)\
        .filter(LibroVenta.id_cliente == cliente_id).all()
    return render_template('ventas_igtf.html', facturas_igtf=facturas_igtf, cliente=cliente)

@main_routes.route('/exportar_libro_ventas_word', methods=['POST'])
def exportar_libro_ventas_word():
    try:
        data = request.get_json()
        html_content = data.get('html_content')
        cliente = data.get('cliente', {})
        os.makedirs('temp', exist_ok=True)
        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(0.8)
            section.right_margin = Cm(0.8)
            section.page_width = Cm(29.7)
            section.page_height = Cm(21.0)
        style = doc.styles['Normal']
        style.font.name = 'Courier New'
        style.font.size = Pt(8)
        # === ENCABEZADO SUPERIOR ===
        header_table = doc.add_table(rows=1, cols=1)
        header_table.autofit = False
        header_table.columns[0].width = Cm(18)
        # Título
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("Libro de Ventas")
        title_run.bold = True
        title_run.font.size = Pt(14)
        doc.add_paragraph()
        # Información de la empresa alineada a la izquierda
        left_cell = header_table.rows[0].cells[0]
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        company_run = left_para.add_run(f"{cliente.get('nombre', 'N/A')}\n")
        company_run.bold = True
        company_run.font.size = Pt(9)
        left_para.add_run(f"R.I.F.  : {cliente.get('rif', 'N/A')}\n")
        left_para.add_run(f"Direccion : {cliente.get('direccion', 'N/A')}\n")
        # Extraer período
        period_text = ""
        company_info = soup.find('div', class_='company-info')
        if company_info:
            for text in company_info.stripped_strings:
                if "Mes :" in text:
                    period_text = text.replace("Mes : ", "")
                    break
        left_para.add_run(f"Mes : {period_text}\n")
        left_para.add_run("Pag:1")
        for row in header_table.rows:
            for cell in row.cells:
                set_cell_border(cell, top={'val': 'nil'}, bottom={'val': 'nil'}, left={'val': 'nil'}, right={'val': 'nil'})
        doc.add_paragraph()
        # === TABLA PRINCIPAL DE DATOS ===
        # Estructura igual a compras pero con columnas de ventas
        main_table = doc.add_table(rows=4, cols=17)
        column_widths = [
            Cm(1.2),  # N° Factura
            Cm(1.2),  # N° Control
            Cm(1.5),  # Fecha Factura
            Cm(2.5),  # RIF
            Cm(3.0),  # Cliente
            Cm(1.5),  # Tipo Documento
            Cm(1.8),  # Monto Total
            Cm(1.5),  # Base Imponible
            Cm(1.2),  # % IVA
            Cm(1.2),  # Monto IVA
            Cm(1.5),  # N° Retención
            Cm(1.2),  # % Retención
            Cm(1.5),  # Monto Retenido
            Cm(1.5),  # Monto IGTF
            Cm(1.2),  # % IGTF
            Cm(1.2),  # Tasa
            Cm(1.2)   # Cant. $
        ]
        for i, width in enumerate(column_widths):
            for row in main_table.rows:
                row.cells[i].width = width
        # === PRIMERA FILA - ENCABEZADOS ===
        row1 = main_table.rows[0]
        headers1 = [
            "N° Factura", "N° Control", "Fecha Factura", "RIF", "Cliente", "Tipo Documento",
            "Monto Total", "Base Imponible", "% IVA", "Monto IVA", "N° Retención", "% Retención",
            "Monto Retenido", "Monto IGTF", "% IGTF", "Tasa", "Cant. $"
        ]
        for i, header in enumerate(headers1):
            row1.cells[i].text = header
            para = row1.cells[i].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
        # === EXTRAER Y AGREGAR DATOS ===
        html_table = soup.find('table')
        totals = [0] * 11  # Ajusta según columnas numéricas
        if html_table:
            for row in html_table.find_all('tr'):
                cells = row.find_all('td')
                if len(cells) >= 17 and not any(cls in row.get('class', []) for cls in ['subtotal', 'total']):
                    data_row = main_table.add_row()
                    cell_data = [cell.get_text().strip() for cell in cells[:17]]
                    for i, data in enumerate(cell_data):
                        if i < len(data_row.cells):
                            cell = data_row.cells[i]
                            cell.text = data
                            para = cell.paragraphs[0]
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 6 else WD_ALIGN_PARAGRAPH.LEFT
                            run = para.runs[0] if para.runs else para.add_run(data)
                            run.font.size = Pt(8)
        # Subtotales y totales (puedes sumar si lo deseas)
        subtotal_row = main_table.add_row()
        subtotal_data = ["" for _ in range(6)] + ["SUB-TOTALES."] + ["" for _ in range(10)]
        for i, data in enumerate(subtotal_data):
            if i < len(subtotal_row.cells):
                cell = subtotal_row.cells[i]
                cell.text = data
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(data)
                run.font.size = Pt(8)
                run.bold = True
                if i >= 6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_row = main_table.add_row()
        total_data = ["" for _ in range(6)] + ["TOTALES."] + ["" for _ in range(10)]
        for i, data in enumerate(total_data):
            if i < len(total_row.cells):
                cell = total_row.cells[i]
                cell.text = data
                para = cell.paragraphs[0]
                run = para.runs[0] if para.runs else para.add_run(data)
                run.font.size = Pt(8)
                run.bold = True
                if i >= 6:
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        border_style = {'val': 'single', 'sz': 4, 'color': '000000'}
        for i, row in enumerate(main_table.rows):
            for j, cell in enumerate(row.cells):
                if i == 0:
                    set_cell_border(cell, top=border_style, bottom=border_style)
                elif i == len(main_table.rows) - 2:
                    set_cell_border(cell, top=border_style)
                elif i == len(main_table.rows) - 1:
                    set_cell_border(cell, bottom=border_style)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        cliente_name = re.sub(r'[^a-zA-Z0-9_-]', '_', cliente.get('nombre', 'cliente'))
        filename = f"libro_ventas_{cliente_name}_{timestamp}.docx"
        filepath = os.path.join('temp', filename)
        doc.save(filepath)
        return jsonify({'filename': filename})
    except Exception as e:
        current_app.logger.error(f'Error al generar Word: {str(e)}')
        return jsonify({'error': f'Error interno: {str(e)}'}, 500)